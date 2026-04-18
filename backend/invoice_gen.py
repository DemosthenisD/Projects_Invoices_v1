"""
Invoice generation: DOCX template filling and PDF conversion.

PDF conversion tries docx2pdf (Microsoft Word) first, then LibreOffice headless as fallback.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document
from shared.config import TEMPLATES_DIR, EXPORTS_DIR


def fill_placeholders(doc: Document, data: dict) -> None:
    """Replace {{KEY}} tokens in all paragraphs and table cells.

    Handles tokens split across multiple runs by consolidating the full
    paragraph text, replacing, then writing back to the first run.
    """
    tokens = {f"{{{{{k}}}}}": str(v) for k, v in data.items()}

    def _replace_in_paragraph(para) -> None:
        full_text = para.text
        if not any(key in full_text for key in tokens):
            return
        if not para.runs:
            return
        new_text = full_text
        for key, value in tokens.items():
            new_text = new_text.replace(key, value)
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""

    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)


def generate_invoice(data: dict, template_name: str, fmt: str = "PDF") -> str:
    """
    Fill the named DOCX template with data and save to exports/.

    Returns the path of the saved DOCX file.
    data keys should match the {{PLACEHOLDER}} names in the template.
    """
    template_path = os.path.join(TEMPLATES_DIR, f"{template_name}.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    fill_placeholders(doc, data)

    os.makedirs(EXPORTS_DIR, exist_ok=True)
    year = data.get("placeholder6", "")
    invoice_no = data.get("placeholder5", "")
    client_code = data.get("placeholder1", "").replace(" ", "_")
    filename = f"{year}_{invoice_no}_{client_code}_Invoice.docx"
    docx_path = os.path.join(EXPORTS_DIR, filename)
    doc.save(docx_path)
    return docx_path


def convert_to_pdf(docx_path: str) -> str:
    """Convert a DOCX file to PDF. Returns the path of the saved PDF file.

    Tries docx2pdf (Microsoft Word) first; falls back to LibreOffice headless.
    Raises RuntimeError if neither converter is available.
    """
    import shutil
    import subprocess

    pdf_path = docx_path.replace(".docx", ".pdf")

    # --- Attempt 1: docx2pdf via Microsoft Word ---
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    # --- Attempt 2: LibreOffice headless ---
    soffice = None
    for candidate in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]:
        if os.path.exists(candidate):
            soffice = candidate
            break
    if not soffice:
        soffice = shutil.which("soffice") or shutil.which("soffice.exe")

    if soffice:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf",
             "--outdir", os.path.dirname(docx_path), docx_path],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            return pdf_path

    raise RuntimeError(
        "PDF conversion failed. Install Microsoft Word (pip install docx2pdf) "
        "or LibreOffice and ensure soffice.exe is reachable."
    )
