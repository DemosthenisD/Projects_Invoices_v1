"""Smoke tests for backend/invoice_gen.py."""
import os
import sys
import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend.invoice_gen import fill_placeholders, generate_invoice


def _make_doc_single_run(text: str) -> Document:
    """One paragraph, one run."""
    doc = Document()
    doc.add_paragraph(text).runs[0].text = text
    # clear the default empty paragraph
    for p in doc.paragraphs[:-1]:
        p._element.getparent().remove(p._element)
    return doc


def _make_doc_split_runs(*parts: str) -> Document:
    """One paragraph with token split across multiple runs."""
    doc = Document()
    para = doc.add_paragraph()
    for part in parts:
        para.add_run(part)
    for p in doc.paragraphs[:-1]:
        p._element.getparent().remove(p._element)
    return doc


def test_single_run_replacement():
    doc = _make_doc_single_run("Hello {{name}}!")
    fill_placeholders(doc, {"name": "World"})
    assert doc.paragraphs[-1].text == "Hello World!"


def test_split_run_replacement():
    """Token split across runs must still be replaced."""
    doc = _make_doc_split_runs("{{place", "holder1}}")
    fill_placeholders(doc, {"placeholder1": "Acme Ltd"})
    assert doc.paragraphs[-1].text == "Acme Ltd"


def test_multiple_tokens_in_paragraph():
    doc = _make_doc_single_run("{{a}} and {{b}}")
    fill_placeholders(doc, {"a": "foo", "b": "bar"})
    assert doc.paragraphs[-1].text == "foo and bar"


def test_no_match_is_unchanged():
    doc = _make_doc_single_run("No placeholders here.")
    fill_placeholders(doc, {"placeholder1": "X"})
    assert doc.paragraphs[-1].text == "No placeholders here."


def test_table_cell_replacement():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("Amount: {{placeholder8}}")
    fill_placeholders(doc, {"placeholder8": "1,000.00"})
    assert table.cell(0, 0).text == "Amount: 1,000.00"


def test_generate_invoice_creates_file(tmp_path, monkeypatch):
    import shared.config as cfg
    monkeypatch.setattr(cfg, "EXPORTS_DIR", str(tmp_path))
    import backend.invoice_gen as ig
    monkeypatch.setattr(ig, "EXPORTS_DIR", str(tmp_path))

    data = {
        "placeholder1": "Test Co",
        "placeholder2": "1 Street",
        "placeholder3": "CY000",
        "placeholder4": "18/04/2026",
        "placeholder5": "99",
        "placeholder6": "2026",
        "placeholder7": "Services",
        "placeholder8": "500.00",
        "placeholder9": "0.00",
        "placeholder8_Exp": "0.00",
        "placeholder9_Exp": "0.00",
        "placeholder8_Tot": "500.00",
        "placeholder9_Tot": "0.00",
        "placeholder10": "500.00",
    }
    path = ig.generate_invoice(data, "template1_v3", "DOCX")
    assert os.path.exists(path)
    assert path.endswith(".docx")

    doc = Document(path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    assert "Test Co" in full_text
    assert "{{" not in full_text, "Unreplaced placeholders remain in output"
